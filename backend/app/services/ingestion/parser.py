"""
Multi-format document parser.
Handles PDFs, DOCX, HTML with OCR fallback for scanned documents.
"""

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

import httpx
from bs4 import BeautifulSoup


@dataclass
class ParsedPage:
    """A parsed page from a document."""

    page_number: int
    content: str
    section_title: Optional[str] = None
    has_tables: bool = False
    has_images: bool = False


@dataclass
class ParsedDocument:
    """A fully parsed document."""

    file_path: str
    file_type: str
    total_pages: int
    pages: List[ParsedPage]
    title: Optional[str] = None
    metadata: Optional[dict] = None


class DocumentParser:
    """Parse documents from various formats into structured text."""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".doc", ".html", ".htm", ".txt"}

    def __init__(self):
        self._section_patterns = [
            r"^(\d+\.)+\s+(.+)$",  # 1.2.3 Section Title
            r"^Section\s+(\d+)[\.:]\s*(.+)$",  # Section 1: Title
            r"^Chapter\s+(\d+)[\.:]\s*(.+)$",  # Chapter 1: Title
            r"^Part\s+(\d+)[\.:]\s*(.+)$",  # Part 1: Title
            r"^([A-Z][A-Z\s]+)$",  # ALL CAPS HEADERS
        ]

    async def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document from file path."""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {suffix}")

        if suffix == ".pdf":
            return await self._parse_pdf(file_path)
        elif suffix in {".docx", ".doc"}:
            return await self._parse_docx(file_path)
        elif suffix in {".html", ".htm"}:
            return await self._parse_html(file_path)
        elif suffix == ".txt":
            return await self._parse_txt(file_path)

        raise ValueError(f"Unhandled format: {suffix}")

    async def parse_from_url(self, url: str) -> ParsedDocument:
        """Parse a document from URL."""
        async with httpx.AsyncClient() as client:
            response = await client.get(url, follow_redirects=True)
            response.raise_for_status()

            content_type = response.headers.get("content-type", "")

            if "pdf" in content_type:
                return await self._parse_pdf_bytes(response.content, url)
            elif "html" in content_type:
                return await self._parse_html_content(response.text, url)
            else:
                raise ValueError(f"Unsupported content type: {content_type}")

    async def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse a PDF file."""
        try:
            import pdfplumber

            pages = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""

                    # Check for tables
                    tables = page.extract_tables()
                    has_tables = len(tables) > 0

                    # Convert tables to text if present
                    if tables:
                        for table in tables:
                            table_text = self._table_to_text(table)
                            if table_text not in text:
                                text += "\n\n" + table_text

                    # Clean the text
                    text = self._clean_text(text)

                    # Detect section title
                    section_title = self._detect_section_title(text)

                    pages.append(
                        ParsedPage(
                            page_number=i + 1,
                            content=text,
                            section_title=section_title,
                            has_tables=has_tables,
                            has_images=len(page.images) > 0,
                        )
                    )

                return ParsedDocument(
                    file_path=file_path,
                    file_type="pdf",
                    total_pages=len(pages),
                    pages=pages,
                    title=self._extract_title(pages),
                )

        except Exception as e:
            # Fallback to PyPDF2 for problematic PDFs
            return await self._parse_pdf_fallback(file_path, str(e))

    async def _parse_pdf_fallback(
        self, file_path: str, error: str
    ) -> ParsedDocument:
        """Fallback PDF parser using PyPDF2."""
        from pypdf import PdfReader

        pages = []
        reader = PdfReader(file_path)

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = self._clean_text(text)
            section_title = self._detect_section_title(text)

            pages.append(
                ParsedPage(
                    page_number=i + 1,
                    content=text,
                    section_title=section_title,
                )
            )

        return ParsedDocument(
            file_path=file_path,
            file_type="pdf",
            total_pages=len(pages),
            pages=pages,
            title=self._extract_title(pages),
            metadata={"parser_fallback": True, "original_error": error},
        )

    async def _parse_pdf_bytes(
        self, content: bytes, source: str
    ) -> ParsedDocument:
        """Parse PDF from bytes."""
        import pdfplumber

        pages = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                text = self._clean_text(text)
                section_title = self._detect_section_title(text)

                pages.append(
                    ParsedPage(
                        page_number=i + 1,
                        content=text,
                        section_title=section_title,
                        has_tables=len(page.extract_tables()) > 0,
                    )
                )

        return ParsedDocument(
            file_path=source,
            file_type="pdf",
            total_pages=len(pages),
            pages=pages,
            title=self._extract_title(pages),
        )

    async def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse a Word document."""
        from docx import Document

        doc = Document(file_path)
        pages = []
        current_page_content = []
        current_section = None
        page_num = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a section header
            if para.style and "Heading" in para.style.name:
                current_section = text

            current_page_content.append(text)

            # Approximate page breaks (every ~3000 chars)
            if len("\n".join(current_page_content)) > 3000:
                pages.append(
                    ParsedPage(
                        page_number=page_num,
                        content="\n\n".join(current_page_content),
                        section_title=current_section,
                    )
                )
                current_page_content = []
                page_num += 1

        # Add remaining content
        if current_page_content:
            pages.append(
                ParsedPage(
                    page_number=page_num,
                    content="\n\n".join(current_page_content),
                    section_title=current_section,
                )
            )

        return ParsedDocument(
            file_path=file_path,
            file_type="docx",
            total_pages=len(pages),
            pages=pages,
            title=self._extract_title(pages),
        )

    async def _parse_html(self, file_path: str) -> ParsedDocument:
        """Parse an HTML file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return await self._parse_html_content(content, file_path)

    async def _parse_html_content(
        self, content: str, source: str
    ) -> ParsedDocument:
        """Parse HTML content."""
        soup = BeautifulSoup(content, "lxml")

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        # Extract title
        title = soup.title.string if soup.title else None

        # Get main content
        main_content = soup.find("main") or soup.find("article") or soup.body or soup

        # Extract text with structure
        text = main_content.get_text(separator="\n", strip=True)
        text = self._clean_text(text)

        # Split into pages (every ~3000 chars for consistency)
        pages = []
        lines = text.split("\n")
        current_content = []
        page_num = 1

        for line in lines:
            current_content.append(line)
            if len("\n".join(current_content)) > 3000:
                pages.append(
                    ParsedPage(
                        page_number=page_num,
                        content="\n".join(current_content),
                        section_title=self._detect_section_title(
                            "\n".join(current_content)
                        ),
                    )
                )
                current_content = []
                page_num += 1

        if current_content:
            pages.append(
                ParsedPage(
                    page_number=page_num,
                    content="\n".join(current_content),
                    section_title=self._detect_section_title(
                        "\n".join(current_content)
                    ),
                )
            )

        return ParsedDocument(
            file_path=source,
            file_type="html",
            total_pages=len(pages),
            pages=pages,
            title=title,
        )

    async def _parse_txt(self, file_path: str) -> ParsedDocument:
        """Parse a plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        content = self._clean_text(content)

        # Split into pages
        pages = []
        lines = content.split("\n")
        current_content = []
        page_num = 1

        for line in lines:
            current_content.append(line)
            if len("\n".join(current_content)) > 3000:
                pages.append(
                    ParsedPage(
                        page_number=page_num,
                        content="\n".join(current_content),
                    )
                )
                current_content = []
                page_num += 1

        if current_content:
            pages.append(
                ParsedPage(
                    page_number=page_num,
                    content="\n".join(current_content),
                )
            )

        return ParsedDocument(
            file_path=file_path,
            file_type="txt",
            total_pages=len(pages),
            pages=pages,
        )

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Replace multiple newlines with double newlines
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Replace multiple spaces with single space
        text = re.sub(r" {2,}", " ", text)

        # Remove non-printable characters except newlines
        text = "".join(
            char for char in text if char.isprintable() or char in "\n\t"
        )

        return text.strip()

    def _detect_section_title(self, text: str) -> Optional[str]:
        """Detect section title from the beginning of text."""
        lines = text.strip().split("\n")
        if not lines:
            return None

        first_line = lines[0].strip()

        for pattern in self._section_patterns:
            match = re.match(pattern, first_line)
            if match:
                return first_line

        return None

    def _extract_title(self, pages: List[ParsedPage]) -> Optional[str]:
        """Extract document title from first page."""
        if not pages or not pages[0].content:
            return None

        first_lines = pages[0].content.split("\n")[:5]
        for line in first_lines:
            line = line.strip()
            if len(line) > 10 and len(line) < 200:
                return line
        return None

    def _table_to_text(self, table: List[List]) -> str:
        """Convert a table to formatted text."""
        if not table:
            return ""

        rows = []
        for row in table:
            cells = [str(cell) if cell else "" for cell in row]
            rows.append(" | ".join(cells))

        return "\n".join(rows)
